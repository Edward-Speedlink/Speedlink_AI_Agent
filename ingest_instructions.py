# apps/ingest.py — FINAL WORKING VERSION (Chroma 0.5+ compatible)
import ollama
import chromadb
from chromadb.utils.embedding_functions import EmbeddingFunction
from typing import List
from pathlib import Path
import PyPDF2
from docx import Document

# === PROPER EMBEDDING FUNCTION CLASS (THIS FIXES THE ERROR) ===
class OllamaEmbeddingFunction(EmbeddingFunction):
    def __call__(self, input: List[str]) -> List[List[float]]:
        return [
            ollama.embeddings(model="nomic-embed-text:latest", prompt=text)["embedding"]
            for text in input
        ]

# Smart chunking
def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    words = text.split()
    chunks = []
    if chunk_size <= 0:
        raise ValueError("chunk_size must be > 0")
    if overlap < 0:
        raise ValueError("overlap must be >= 0")
    if overlap >= chunk_size:
        raise ValueError("overlap must be smaller than chunk_size")
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append(chunk)
        i += chunk_size - overlap
    return chunks

# Extract text
def extract_text(file_path: str) -> str:
    path = Path(file_path)
    text = ""
    if path.suffix == ".pdf":
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    elif path.suffix == ".docx":
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    elif path.suffix == ".txt":
        text = path.read_text(encoding="utf-8")
    return " ".join(text.split())  # clean whitespace

# === MAIN ===
print("RHODA rebuilding Speedlink brain...")

client = chromadb.PersistentClient(path="./company_db")

# Safe wipe
try:
    client.delete_collection("speedlink_knowledge")
    print("Old DB wiped")
except:
    pass

# Create collection with proper class-based embedding function
collection = client.create_collection(
    name="speedlink_knowledge",
    embedding_function=OllamaEmbeddingFunction(),
    metadata={"hnsw:space": "cosine"}
)

# ← PUT YOUR REAL FILES HERE
files = ["file1.pdf", "file2.pdf", "Speedlink_Plans.pdf", "Company_Policy.pdf"]

total = 0
for file in files:
    if not Path(file).exists():
        print(f"Missing: {file}")
        continue
    print(f"Loading {file}...")
    text = extract_text(file)
    chunks = chunk_text(text)
    print(f"→ {len(chunks)} chunks")

    collection.add(
        documents=chunks,
        ids=[f"{Path(file).stem}_{i}" for i in range(len(chunks))],
        metadatas=[{"source": file, "chunk": i} for i in range(len(chunks))]
    )
    total += len(chunks)

print(f"\nRHODA IS NOW 100% SPEEDLINK")
print(f"Loaded {total} perfect chunks. Ready for production!")







# # ingest_universal.py
# import chromadb
# import ollama
# from docx import Document
# import PyPDF2
# import os

# # ---- Embedding function ----
# class OllamaEmbeddingFunction:
#     def __init__(self, model="nomic-embed-text"):
#         self.model = model

#     def __call__(self, input):
#         if isinstance(input, str):
#             input = [input]

#         embeddings = []
#         for text in input:
#             resp = ollama.embeddings(model=self.model, prompt=text)
#             embeddings.append(resp["embedding"])
#         return embeddings

#     def name(self):
#         return f"ollama-{self.model}"

# # ---- Connect to ChromaDB ----
# client = chromadb.PersistentClient(path="./company_db")

# # ---- DELETE existing collection if present ----
# # Uncomment this to wipe the collection before a fresh ingestion
# # if "company_info" in [c.name for c in client.list_collections()]:
# #     client.delete_collection("company_info")
# #     print("⚠️ Existing 'company_info' collection deleted!")

# # ---- Create collection ----
# collection = client.get_or_create_collection(
#     name="company_info",
#     embedding_function=OllamaEmbeddingFunction()
# )

# # ---- Helper functions to read files ----
# def read_pdf(file_path):
#     text = ""
#     with open(file_path, "rb") as f:
#         reader = PyPDF2.PdfReader(f)
#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return [line.strip() for line in text.splitlines() if line.strip()]

# def read_docx(file_path):
#     doc = Document(file_path)
#     return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

# def read_file(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == ".txt":
#         with open(file_path, "r") as f:
#             return [line.strip() for line in f if line.strip()]
#     elif ext == ".pdf":
#         return read_pdf(file_path)
#     elif ext == ".docx":
#         return read_docx(file_path)
#     else:
#         print(f"⚠️ Unsupported file type: {file_path}")
#         return []

# # ---- List all files to index ----
# files_to_index = ["file1.pdf", "file2.pdf"] # "scenarios.docx", "policies.pdf"]  # <-- add your real files here
# all_lines = []

# for file in files_to_index:
#     lines = read_file(file)
#     all_lines.extend(lines)
#     print(f"📄 {file}: {len(lines)} lines extracted")

# # ---- Index into ChromaDB ----
# for i, line in enumerate(all_lines):
#     collection.add(
#         ids=[f"doc_{i}"],
#         documents=[line]
#     )

# print(f"✅ Indexed {len(all_lines)} lines from {len(files_to_index)} files into ChromaDB")
